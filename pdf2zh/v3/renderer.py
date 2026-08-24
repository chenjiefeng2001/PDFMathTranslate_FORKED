"""Module: Renderer — V4 Epic C.
Multi-format rendering from VisualTree.

Phase 2 upgrades:
  - SVGRenderer (vector SVG output)
  - DOCXRenderer (basic DOCX via python-docx)

Usage:
    from pdf2zh.v3.renderer import PDFRenderer, HTMLRenderer, SVGRenderer
    renderer = PDFRenderer()
    output = renderer.render(visual_tree)
"""

from __future__ import annotations
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple
from pdf2zh.v3.visual_tree import (
    VisualTree,
    VisualNode,
    VisualNodeType,
    BoundingBox,
    Page,
    Paragraph,
    Line,
    TextRun,
)

logger = logging.getLogger(__name__)


@dataclass
class RenderContext:
    fmt: str = "pdf"
    dpi: int = 300
    font_scale: float = 1.0
    options: dict = field(default_factory=dict)


class Renderer(ABC):
    @abstractmethod
    def render(
        self, tree: VisualTree, context: Optional[RenderContext] = None
    ) -> bytes:
        pass

    @abstractmethod
    def render_page(self, page: Page, context: Optional[RenderContext] = None) -> bytes:
        pass


class PDFRenderer(Renderer):
    def render(self, tree: VisualTree, context=None):
        ctx = context or RenderContext()
        output = []
        for page in tree.pages:
            output.append(self.render_page(page, ctx))
        return b"\n".join(output)

    def render_page(self, page: Page, context=None):
        ctx = context or RenderContext()
        lines = [
            f"%PDF-page num={page.page_num} w={page.width:.0f} h={page.height:.0f}"
        ]
        for node in page.walk():
            if isinstance(node, TextRun):
                lines.append(
                    f"  text '{node.text[:40]}' @({node.bbox.x:.0f},{node.bbox.y:.0f}) font={node.font or 'default'} {len(node.text)}ch"
                )
            elif isinstance(node, Line):
                lines.append(
                    f"  line y={node.bbox.y:.0f} baseline={node.baseline:.0f} alignment={node.alignment or 'left'}"
                )
            elif isinstance(node, Paragraph):
                lines.append(
                    f"  para id={node.id} ({node.bbox.x:.0f},{node.bbox.y:.0f},{node.bbox.width:.0f}x{node.bbox.height:.0f}) lang={node.language or ''}"
                )
        return ("\n".join(lines)).encode("utf-8")


class HTMLRenderer(Renderer):
    def render(self, tree: VisualTree, context=None):
        ctx = context or RenderContext()
        parts = [
            "<!DOCTYPE html><html><head><meta charset='utf-8'><style>"
            "body{font-family:serif;margin:2em}"
            ".page{page-break-after:always;margin-bottom:2em;border:1px solid #ccc;padding:1em}"
            ".para{margin-bottom:0.5em}.line{white-space:pre-wrap}"
            "</style></head><body>"
        ]
        for page in tree.pages:
            parts.append(f"<div class='page'><h3>Page {page.page_num + 1}</h3>")
            for node in page.walk():
                if isinstance(node, Paragraph):
                    parts.append(f"<div class='para' id='{node.id}'>")
                elif isinstance(node, Line):
                    parts.append(f"<span class='line'>")
                elif isinstance(node, TextRun):
                    parts.append(node.text)
                elif isinstance(node, Line):
                    parts.append("</span>")
                elif isinstance(node, Paragraph):
                    parts.append("</div>")
            parts.append("</div>")
        parts.append("</body></html>")
        return "\n".join(parts).encode("utf-8")


class SVGRenderer(Renderer):
    """Render VisualTree to SVG format."""

    def render(self, tree, context=None):
        ctx = context or RenderContext()
        parts = ['<?xml version="1.0" encoding="utf-8"?>']
        parts.append('<svg xmlns="http://www.w3.org/2000/svg">')
        for page in tree.pages:
            parts.append(
                f'<g id="page-{page.page_num}" transform="translate(0,{page.page_num * (page.height + 20)})">'
            )
            parts.append(
                f'<rect x="0" y="0" width="{page.width}" height="{page.height}" fill="white" stroke="#ccc"/>'
            )
            for node in page.walk():
                if isinstance(node, TextRun):
                    x = node.bbox.x if hasattr(node, "bbox") else 0
                    y = node.bbox.y if hasattr(node, "bbox") else 0
                    fs = (
                        node.font_size
                        if hasattr(node, "font_size") and node.font_size
                        else 12
                    )
                    safe = (
                        node.text.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    parts.append(
                        f'<text x="{x:.1f}" y="{y + fs:.1f}" font-size="{fs}">{safe}</text>'
                    )
            parts.append("</g>")
        parts.append("</svg>")
        return "\n".join(parts).encode("utf-8")

    def render_page(self, page, context=None):
        ctx = context or RenderContext()
        parts = ['<?xml version="1.0" encoding="utf-8"?>']
        parts.append(
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{page.width}" height="{page.height}">'
        )
        parts.append(
            f'<rect x="0" y="0" width="{page.width}" height="{page.height}" fill="white"/>'
        )
        for node in page.walk():
            if isinstance(node, TextRun):
                x = node.bbox.x if hasattr(node, "bbox") else 0
                y = node.bbox.y if hasattr(node, "bbox") else 0
                fs = (
                    node.font_size
                    if hasattr(node, "font_size") and node.font_size
                    else 12
                )
                safe = (
                    node.text.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                )
                parts.append(
                    f'<text x="{x:.1f}" y="{y + fs:.1f}" font-size="{fs}">{safe}</text>'
                )
        parts.append("</svg>")
        return "\n".join(parts).encode("utf-8")


class MarkdownRenderer(Renderer):
    def render(self, tree: VisualTree, context=None):
        lines = []
        for page in tree.pages:
            lines.append(f"\n## Page {page.page_num + 1}\n")
            for node in page.walk():
                if isinstance(node, TextRun):
                    lines.append(node.text)
                elif isinstance(node, Line):
                    lines.append("  \n")
        return ("".join(lines)).encode("utf-8")

    def render_page(self, page: Page, context=None):
        return self.render_page_as_markdown(page)

    def render_page_as_markdown(self, page: Page) -> str:
        lines = []
        lines.append(f"## Page {page.page_num + 1}\n")
        for node in page.walk():
            if isinstance(node, TextRun):
                lines.append(node.text)
            elif isinstance(node, Line):
                lines.append("  \n")
        return ("".join(lines)).encode("utf-8")


class DOCXRenderer(Renderer):
    """Render VisualTree to DOCX format (basic).
    Requires python-docx package.
    """

    def render(self, tree, context=None):
        ctx = context or RenderContext()
        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            logger.warning("python-docx not installed; falling back to plain text")
            md = MarkdownRenderer()
            return md.render(tree, context)
        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(12)
        for page in tree.pages:
            for node in page.walk():
                if isinstance(node, Paragraph):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                elif isinstance(node, TextRun):
                    p = (
                        doc.add_paragraph()
                        if not hasattr(node, "_added_p")
                        else doc.paragraphs[-1]
                    )
                    run = p.add_run(node.text)
                    if node.font:
                        run.font.name = node.font
        import io

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def render_page(self, page, context=None):
        tree = VisualTree()
        tree.add_page(page)
        return self.render(tree, context)


class RendererFactory:
    _registry = {
        "pdf": PDFRenderer,
        "html": HTMLRenderer,
        "markdown": MarkdownRenderer,
        "svg": SVGRenderer,
        "docx": DOCXRenderer,
    }

    @classmethod
    def create(cls, fmt: str) -> Renderer:
        renderer_cls = cls._registry.get(fmt.lower())
        if renderer_cls is None:
            raise ValueError(
                f"Unsupported format: {fmt}. Supported: {list(cls._registry.keys())}"
            )
        return renderer_cls()

    @classmethod
    def register(cls, fmt: str, renderer_cls) -> None:
        cls._registry[fmt] = renderer_cls

    @classmethod
    def supported_formats(cls) -> List[str]:
        return list(cls._registry.keys())


__all__ = [
    "RenderContext",
    "Renderer",
    "PDFRenderer",
    "HTMLRenderer",
    "MarkdownRenderer",
    "SVGRenderer",
    "DOCXRenderer",
    "RendererFactory",
]
